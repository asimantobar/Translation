# -*- coding: utf-8 -*-

import Levenshtein
from docx import Document
import os
import pandas as pd
import re
import random
from mtranslate import translate
import shutil


def extract_text(element):
    if hasattr(element, 'text'):
        return element.text.strip()
    elif element._element is not None:
        return ' '.join(extract_text(e) for e in element._element)

def content_docx(docx_file):
    doc = Document(docx_file)
    content = []

    for section in doc.sections:
        for header in section.header.paragraphs:
            content.append(extract_text(header))
            
    for paragraph in doc.paragraphs:
        content.append(extract_text(paragraph))
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                content.append(extract_text(cell))
    return content
    
def merge_content(folder_path, output_file): 
    df = pd.DataFrame(columns=['Filename', 'Content'])
    cont = 0 
    with open(output_file, 'w', encoding="utf-8") as txt_file:
        for filename in os.listdir(folder_path):
            if filename.endswith(".docx"):
                docx_path = os.path.join(folder_path, filename)
                content = content_docx(docx_path)
                txt_file.write('\n'.join(content))
                for line in content:
                    df.loc[cont] = [filename, line]
                    cont += 1
    
    df.dropna(subset=['Content'], inplace=True)  
    df = df[df['Content'] != ''] 
    df.reset_index(drop=True, inplace=True) 
    return df

def generar_codigo(usados):
    code = random.randint(1000, 9999)
    while code in usados:
        code = random.randint(1000, 9999)
    return code

def extraer_palab_cod(df, usados):
    selected_words = []
    for text in df['Content']:
        words = []
        split_text = text.split()
        for word in split_text:
            if word.startswith("ID") and "_" in word:
                if word not in words:
                    words.append(word)
            elif word.startswith("@") and word.endswith("@") and len(word) > 1:
                if word not in words:
                    words.append(word)
        for word in words:
            if word not in selected_words:
                code = generar_codigo(usados)
                selected_words.append({'word': word, 'code': code})
                usados.add(code)
    selected_df = pd.DataFrame(selected_words)
    return selected_df


def sustituir_cod_asociado(df, codes_df):
    df3 = df.copy()
    for index, row in codes_df.iterrows():
        word = row['word']
        code = row['code']
        df3['Content'] = df3['Content'].str.replace(re.escape(word), str(code))
    return df3


def extract_var2(df, used_codes):
    var2 = pd.DataFrame(columns=['word', 'code'])

    pattern = r'\b(\d{4})_(\w+)\b'

    for text in df['Content']:
        matches = re.findall(pattern, text)
        for match in matches:
            code = generar_codigo(used_codes)
            word = match[0] + '_' + match[1]
            var2 = var2.append({'word': word, 'code': code}, ignore_index=True)
            used_codes.add(code)

        words = re.findall(r'\b(\d{4})(\S+)\b', text)
        for word in words:
            code = generar_codigo(used_codes)
            word = word[0] + word[1]
            var2 = var2.append({'word': word, 'code': code}, ignore_index=True)
            used_codes.add(code)

    return var2

def sustituir_palabra_asociada(df, codes_df):
    for column in df.columns:
        for index, row in codes_df.iterrows():
            word = row['word']
            code = row['code']
            df[column] = df[column].str.replace(str(code), re.escape(word))
    return df

def compare_strings(strings):
    num_strings = len(strings)
    diferencias = [[0] * num_strings for _ in range(num_strings)]

    for i in range(num_strings):
        for j in range(i + 1, num_strings):
            dif = Levenshtein.distance(strings[i], strings[j])
            diferencias[i][j] = diferencias[j][i] = dif

    return diferencias


def translate_files(args):
        
    folder_path = args.folder_path
    name = args.name

    print(f"The {args.folder_path} folder has been chosen to translate it's documents")    

    translations_folder = os.path.join(folder_path, "Translations")
    os.makedirs(translations_folder, exist_ok=True)
    
    output_file = translations_folder + '/Plantillas_EN_' + name + '.txt'
        
    if not folder_path:
        print("Error: No se ha seleccionado ningun folder. Por Favor, seleccione carpeta con documentos para traducir")
        return

    if not name:
        name = os.path.basename(folder_path)
        
    print(f" The documents related to the translation can be easily identified under the {name}") 
    
    df = merge_content(folder_path, output_file)
    
    df2 = df.drop_duplicates(subset='Content')
    df2 = df.groupby('Content')['Filename'].apply(set).apply(', '.join).reset_index(name='Filename')

    usados = set()
    variables_df = extraer_palab_cod(df2, usados)
    
    df3 = sustituir_cod_asociado(df2,variables_df)
    var2 = extract_var2(df3, usados)
    df4 = sustituir_cod_asociado(df3,var2)

    matches = [r'^\d+\s*/?\s*\d+$', r'^\d+\s*º$|\d+\s*%$|\d+°$', r'^\s*\d+\s*$', r'^(?:[<>]=?)?\d+º=\d+$']
    contain = [r'\b\d+\s*%\s*–\s*\d+\b', r'^\s*[-\[\]]\s*$']
    
    df4 = df4[~df4['Content'].str.isnumeric()] 
    
    for m in range(len(matches)):
        df4 = df4[~df4['Content'].str.match(matches[m])]  
    
    for c in range(len(contain)):
        df4 = df4[~df4['Content'].str.contains(contain[c], regex=False, na=False)]   
    
    df4 = df4.reset_index(drop=True)


    target_lang = {}
    
    if args.ES:
        target_lang.update({'Translation_ES' : 'es'})
        print("The documents will be translated to: Spanish") 
    if args.ZH:
        target_lang.update({'Translation_ZH' : 'zh-CN'})
        print("The documents will be translated to: Chinese (simplified)") 
    if args.DE:
        target_lang.update({'Translation_DE': 'de'})
        print("The documents will be translated to: German") 
    if args.FR:
        target_lang.update({'Translation_FR': 'fr'})
        print("The documents will be translated to: French") 
    if args.PT:
        target_lang.update({'Translation_PT': 'pt'})
        print("The documents will be translated to: Portuguese") 
    if args.IT:
        target_lang.update({'Translation_IT': 'it'})
        print("The documents will be translated to: Italian") 
    if args.EU:
        target_lang.update({'Translation_EU': 'eu'})
        print("The documents will be translated to: Basque") 
    if args.JA:
        target_lang.update({'Translation_JA': 'ja'})
        print("The documents will be translated to: Japanese") 
    if args.AR:
        target_lang.update({'Translation_AR': 'ar'})
        print("The documents will be translated to: Arabic") 
        

   
    for column, language in target_lang.items():
        translations = []
        print(f"The translation to {language} is currently being realized") 
        for content in df4['Content']:
            try:
                translation = translate(content,  language,'en')  
                translations.append(translation)
            except Exception as e:
                print(f"Error occurred during translation: {str(e)}")
                translations.append(None)  
    
        df4[column] = translations

    df4 = sustituir_palabra_asociada(df4, var2)
    df4 = sustituir_palabra_asociada(df4, variables_df)
   
    content_strings = df4['Content'].tolist()
    diferencias = compare_strings(content_strings)
    
    output =  translations_folder + '/difference_units_'+ name +'.txt'

    with open(output, 'w', encoding="utf-8") as file:
        print("The string comparison to filter similar phrases is being conducted") 
        for i in range(len(content_strings)):
            for j in range(i + 1, len(content_strings)):
                ignorar = False
    
                if any(word.startswith("ID_") for word in content_strings[i].split()): 
                    ignorar = True
                if len(content_strings[i].split()) == 1 and len(content_strings[j].split()) == 1:  
                    ignorar = True
                
                if content_strings[i].lower().startswith("max") and content_strings[j].lower().startswith("min"):
                    ignorar = True
                elif content_strings[i].lower().startswith("min") and content_strings[j].lower().startswith("max"):
                    ignorar = True
                
                if content_strings[i].lower() == "ankle" and content_strings[j].lower() == "knee":
                    ignorar = True
                elif content_strings[i].lower() == "knee" and content_strings[j].lower() == "ankle":
                    ignorar = True
                
                if content_strings[i].lower() == "external" and content_strings[j].lower() == "internal":
                    ignorar = True
                elif content_strings[i].lower() == "internal" and content_strings[j].lower() == "external":
                    ignorar = True
                    
                if content_strings[i].lower() == "lower" and content_strings[j].lower() == "upper":
                    ignorar = True
                elif content_strings[i].lower() == "upper" and content_strings[j].lower() == "lower":
                    ignorar = True
    
                
                if not ignorar and diferencias[i][j] <= 3:
                    file.write(f"Strings {i} y {j} tienen {diferencias[i][j]} caracteres diferentes.\n")
                    file.write(f'String {i}: {content_strings[i]} \n    En documentos: '+df4['Filename'].loc[i]+'\n')
                    file.write(f'String {j}: {content_strings[j]} \n    En documentos: '+ df4['Filename'].loc[j] + '\n\n')
                    

    catalogpath = translations_folder + '/Plantillas(PO)_'+ name +'.txt'

    header = [
        '#',
        'msgid ""',
        'msgstr ""',
        '"Project-Id-Version: Plantillas\\n"',
        '"POT-Creation-Date: 2023-06-16 12:02+0100\\n"',
        '"PO-Revision-Date: 2023-06-16 12:03+0100\\n"',
        '"Last-Translator: Ana Siman <anabeatrizsiman@gmail.com>\\n"',
        '"Language-Team: STT Systems\\n"',
        '"MIME-Version: 1.0\\n"',
        '"Content-Type: text/plain; charset=UTF-8\\n"',
        '"Content-Transfer-Encoding: 8bit\\n"',
        '"X-Poedit-Basepath: .\\n"',
        '"X-Poedit-Country: Spain\\n"',
        '"X-Poedit-KeywordsList: _;gettext;gettext_noop\\n"',
        '"X-Poedit-Language: English\\n"',
        ''
    ]
    
    everyword = ['# ', 'msgid', 'msgstr']
    
    with open(catalogpath, 'w', encoding="utf-8") as po_file:
        for line in header:
            po_file.write(line)
            po_file.write('\n')
            
        for i in range(len(df4)):
            po_file.write(everyword[0] + df4["Filename"].values[i] + '\n')
            po_file.write(everyword[1] + ' "' + 
                          df4["Content"].values[i].replace('\n', ' ') + '"\n')
            po_file.write(everyword[2] + ' ""\n\n')
   
    idiomas = {}
    
    if args.ES:
        idiomas.update({'ES': '/Plantillas_ES_'})
    if args.ZH:
        idiomas.update({'ZH': '/Plantillas_ZH_'})
    if args.DE:
        idiomas.update({'DE': '/Plantillas_DE_'})
    if args.FR:
        idiomas.update({'FR': '/Plantillas_FR_'})
    if args.PT:
        idiomas.update({'PT': '/Plantillas_PT_'})
    if args.IT:
        idiomas.update({'IT': '/Plantillas_IT_'})
    if args.EU:
        idiomas.update({'EU': '/Plantillas_EU_'})
    if args.JA:
        idiomas.update({'JA': '/Plantillas_JA_'})
    if args.AR:
        idiomas.update({'AR': '/Plantillas_AR_'})        
        
    for language, path in idiomas.items():
        print(f"The .po file is being generated for the {language} translation") 
        translations = df4[f"Translation_{language}"]
        filepath = translations_folder +'/'+ path + name + '.txt'
        
        with open(filepath, 'w', encoding="utf-8") as po_file:
            for line in header:
                po_file.write(line)
                po_file.write('\n')
            for i in range(len(df4)):
                po_file.write(everyword[0] + df4["Filename"].values[i] + '\n')
                po_file.write(everyword[1] + ' "' + 
                              df4["Content"].values[i].replace('\n', ' ') + '"\n')
                po_file.write(everyword[2] + ' "' + 
                              translations.values[i].replace('\n', ' ') + '"\n\n')

        target = translations_folder + '/ Translation_'+ name +'_' + language + '.po'
        shutil.copyfile(filepath, target)
